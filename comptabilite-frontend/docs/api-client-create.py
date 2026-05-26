"""
Run this script once to generate api-client-create.docx
Requires: pip install python-docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Styles helpers ────────────────────────────────────────────────────────────
def set_heading(para, text, level=1):
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(16 - (level - 1) * 2)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F3F4F6")
    p._p.pPr.append(shading) if p._p.pPr is not None else None
    return p

def add_table_row(table, col1, col2, header=False):
    row = table.add_row()
    for i, text in enumerate([col1, col2]):
        cell = row.cells[i]
        cell.text = text
        run = cell.paragraphs[0].runs[0]
        run.bold = header
        run.font.size = Pt(10)

# ── Title ─────────────────────────────────────────────────────────────────────
title = doc.add_heading("", 0)
set_heading(title, "API Documentation — POST /api/clients", level=1)
doc.add_paragraph("Version 1.0  |  Date : 2026-05-19  |  Projet : Facturation")
doc.add_paragraph()

# ── 1. Overview ───────────────────────────────────────────────────────────────
h = doc.add_paragraph()
set_heading(h, "1. Vue d'ensemble", level=2)

doc.add_paragraph("Crée un nouveau profil client dans le système de facturation.")

overview = doc.add_table(rows=1, cols=2)
overview.style = "Table Grid"
add_table_row(overview, "Attribut", "Valeur", header=True)
for k, v in [
    ("Méthode", "POST"),
    ("URL", "/api/clients"),
    ("Auth", "Bearer JWT (header Authorization)"),
    ("Content-Type", "application/json"),
]:
    add_table_row(overview, k, v)
doc.add_paragraph()

# ── 2. Request Body ───────────────────────────────────────────────────────────
h = doc.add_paragraph()
set_heading(h, "2. Corps de la requête", level=2)

add_code_block(doc, """{
  "companyName": "Acme Corp International",      // obligatoire
  "website": "https://www.acme.com",

  "contact": {
    "fullName": "John Doe",
    "email": "john.doe@acme.com",                // obligatoire
    "phone": "+33 6 12 34 56 78"
  },

  "billingAddress": {
    "street": "123 Financial District, Suite 400",
    "city": "Paris",
    "postalCode": "75001",
    "country": "France"
  },

  "financial": {
    "taxId": "FR 12 345 678 901",
    "currency": "EUR",
    "paymentTerms": "Net 30"
  }
}""")
doc.add_paragraph()

# champs obligatoires
doc.add_paragraph("Champs obligatoires : companyName, contact.email").runs[0].bold = True

# valeurs acceptées
doc.add_paragraph()
h = doc.add_paragraph()
set_heading(h, "Valeurs acceptées", level=3)

t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
add_table_row(t, "Champ", "Valeurs", header=True)
add_table_row(t, "currency", "EUR | GBP | USD | CHF")
add_table_row(t, "paymentTerms", "Net 15 | Net 30 | Net 45 | Net 60 | Immédiat")
doc.add_paragraph()

# ── 3. Responses ──────────────────────────────────────────────────────────────
h = doc.add_paragraph()
set_heading(h, "3. Réponses", level=2)

for code, title_, body in [
    ("201 Created", "Succès", """{
  "id": "clnt_01j9z...",
  "reference": "CUST-2026-4782",
  "companyName": "Acme Corp International",
  "contact": { "fullName": "John Doe", "email": "john.doe@acme.com", "phone": "+33 6 12 34 56 78" },
  "billingAddress": { "street": "...", "city": "Paris", "postalCode": "75001", "country": "France" },
  "financial": { "taxId": "FR 12 345 678 901", "currency": "EUR", "paymentTerms": "Net 30" },
  "createdAt": "2026-05-19T10:30:00Z"
}"""),
    ("400 Bad Request", "Erreur de validation", """{
  "error": "VALIDATION_ERROR",
  "details": [
    { "field": "companyName", "message": "required" },
    { "field": "contact.email", "message": "invalid email format" }
  ]
}"""),
    ("401 Unauthorized", "Token absent ou expiré", '{ "error": "UNAUTHORIZED" }'),
    ("409 Conflict", "Client avec ce taxId déjà existant", """{
  "error": "DUPLICATE_CLIENT",
  "message": "A client with taxId 'FR 12 345 678 901' already exists",
  "existingId": "clnt_00x4..."
}"""),
]:
    p = doc.add_paragraph()
    run = p.add_run(f"{code} — {title_}")
    run.bold = True
    run.font.size = Pt(11)
    add_code_block(doc, body)
    doc.add_paragraph()

# ── 4. Payment Terms ──────────────────────────────────────────────────────────
h = doc.add_paragraph()
set_heading(h, "4. Utilité du champ paymentTerms", level=2)

doc.add_paragraph(
    "paymentTerms définit le délai accordé au client pour payer une facture après sa date d'émission. "
    "Il s'applique par défaut à toutes les factures créées pour ce client."
)
doc.add_paragraph()

t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
add_table_row(t, "Valeur", "Signification", header=True)
for v, s in [
    ("Net 15", "Paiement sous 15 jours"),
    ("Net 30", "Paiement sous 30 jours (standard B2B France)"),
    ("Net 45", "Paiement sous 45 jours"),
    ("Net 60", "Paiement sous 60 jours (maximum légal — LME 2008)"),
    ("Immédiat", "Paiement dû à réception de la facture"),
]:
    add_table_row(t, v, s)
doc.add_paragraph()

doc.add_paragraph(
    "Lors de la création d'une facture, ce champ calcule automatiquement la date d'échéance :"
)
add_code_block(doc, "dueDate = issueDate + paymentTerms (jours)\n\n"
               "Exemple : émission le 2026-05-19, Net 30  →  dueDate = 2026-06-18")
doc.add_paragraph()
doc.add_paragraph("Ce champ pilote également :")
for item in [
    "le statut de la facture (en_retard si dueDate dépassée)",
    "les relances automatiques (J-3, J+1, J+7)",
    "les KPIs du dashboard (Pending, Overdue)",
]:
    doc.add_paragraph(item, style="List Bullet")

# ── 5. Notes d'implémentation ─────────────────────────────────────────────────
doc.add_paragraph()
h = doc.add_paragraph()
set_heading(h, "5. Notes d'implémentation", level=2)

for item in [
    "La référence (CUST-2026-XXXX) doit être générée côté serveur — ne pas utiliser la valeur du front.",
    "L'id retourné devra remplacer clientName dans le modèle Invoice (liaison client ↔ facture).",
    "La méthode save() dans NewClientComponent doit appeler POST /api/clients et gérer les erreurs 400/409.",
]:
    doc.add_paragraph(item, style="List Bullet")

# ── Save ──────────────────────────────────────────────────────────────────────
import os
out = os.path.join(os.path.dirname(__file__), "api-client-create.docx")
doc.save(out)
print(f"Document généré : {out}")
